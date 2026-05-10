from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from src.common.utils import project_root


REPO_URL = "https://github.com/lhrscw/deep-learning-spatial-intelligence-hw2"
WEIGHTS_URL = f"{REPO_URL}/releases/download/v1.0.0/hw2_weights.zip"
TRACKING_URL = f"{REPO_URL}/releases/download/v1.0.0/hw2_tracking_demo.zip"


def set_run_font(run, size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_doc_fonts(doc: Document) -> None:
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    doc.styles["Normal"].font.size = Pt(10.5)


def clear_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def add_paragraph(doc: Document, text: str = "", style: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run_font(run)
    paragraph.paragraph_format.first_line_indent = Cm(0.74) if style is None else None
    paragraph.paragraph_format.line_spacing = 1.25
    return paragraph


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, size=9)
    run.italic = True


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(level=level)
    run = paragraph.add_run(text)
    set_run_font(run, size=16 if level == 1 else 13, bold=True)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], caption: str | None = None) -> None:
    if caption:
        add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run_font(run, size=9)


def add_picture_if_exists(doc: Document, rel_path: str, caption: str, width_cm: float = 15.5) -> None:
    path = project_root() / rel_path
    if not path.exists():
        add_paragraph(doc, f"图像文件暂缺：{rel_path}")
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def new_document(template: Path | None) -> Document:
    if template and template.exists():
        doc = Document(template)
        clear_body(doc)
    else:
        doc = Document()
    set_doc_fonts(doc)
    section = doc.sections[0]
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    return doc


def add_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("深度学习与空间智能 HW2 实验报告")
    set_run_font(run, size=18, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Flowers102 分类、Road Vehicle 检测跟踪、U-Net 语义分割")
    set_run_font(run, size=12)

    add_paragraph(doc, "小组成员：待填写")
    add_paragraph(doc, "学号：待填写")
    add_paragraph(doc, "分工：待填写")
    add_paragraph(doc, f"代码仓库链接：{REPO_URL}")
    add_paragraph(doc, f"模型权重下载地址：{WEIGHTS_URL}")
    add_paragraph(doc, f"跟踪演示视频与日志：{TRACKING_URL}")


def build_report(output: Path, template: Path | None = None) -> None:
    doc = new_document(template)
    add_cover(doc)

    add_heading(doc, "摘要")
    add_paragraph(
        doc,
        "本报告围绕 HW2 的三个任务展开，分别对应图像分类、目标检测跟踪和语义分割三类典型空间智能问题。任务一在 Flowers102 数据集上微调 ImageNet 预训练 ResNet-18，并通过随机初始化、学习率组合和 SE 注意力模块进行对比；任务二在 Road Vehicle 数据集上训练 YOLOv8n，进一步使用多目标跟踪得到稳定 Tracking ID，并基于虚拟线完成车辆越线计数；任务三从零实现 U-Net，在 Stanford Background Dataset 上比较 Cross-Entropy、Dice Loss 和 CE+Dice 三种损失函数。",
    )
    add_paragraph(
        doc,
        "实验结果显示，预训练特征对小规模细粒度分类非常关键，ResNet-18 随机初始化在 8 epoch 内只有 0.2340 的测试准确率，而 ImageNet 预训练加 SE-block 可达到 0.8676。YOLOv8n 在 20 epoch 设置下取得 mAP50 0.3770，并在 12 秒演示视频中产生 67 个唯一 Tracking ID 和 9 次越线计数。U-Net 分割实验中，Cross-Entropy 在 10 epoch 设置下取得最高验证 mIoU 0.5394，说明在该数据和训练长度下，像素级监督比单独 Dice Loss 更稳定。",
    )
    add_paragraph(
        doc,
        "本报告使用当前可复现实验结果形成 Word 初稿，保留完整训练命令、公开代码仓库、模型权重和演示视频链接。后续可在 Word 中补充小组信息、调整版式，并根据更长训练或实拍道路视频替换对应结果。",
    )
    add_paragraph(doc, "关键词：Flowers102；ResNet；YOLOv8；多目标跟踪；U-Net；Dice Loss；空间智能")

    add_table(
        doc,
        ["项目", "内容"],
        [
            ["代码仓库", REPO_URL],
            ["模型权重", WEIGHTS_URL],
            ["跟踪演示", TRACKING_URL],
            ["报告 Word", "reports/HW2_report_draft.docx"],
            ["结果汇总", "outputs/logs/"],
        ],
        "表1 实验资源与公开链接",
    )

    add_heading(doc, "一、研究背景与总体目标")
    add_paragraph(
        doc,
        "空间智能任务通常需要模型理解图像中的目标类别、空间位置、运动关系和像素级区域。单一分类准确率无法覆盖这些能力，因此本次作业把实验拆成三个层级：第一层是图像级分类，判断一张花卉图像属于哪一类；第二层是目标级检测与跟踪，在道路场景中定位车辆并维护跨帧身份；第三层是像素级分割，对场景中的每个像素预测语义类别。",
    )
    add_paragraph(
        doc,
        "三个任务在建模方式上形成递进关系。ResNet 分类任务关注全局判别特征，适合分析预训练、学习率和注意力模块的影响；YOLOv8 检测任务关注目标框、置信度和类别预测，并通过跟踪算法把静态检测结果扩展到时间维度；U-Net 分割任务则要求编码器和解码器在局部纹理与全局上下文之间建立联系，损失函数设计会直接影响边界、细小区域和类别不平衡。",
    )
    add_paragraph(
        doc,
        "为了保证结果可复现，所有训练命令、配置文件、日志摘要、图表和报告草稿均保留在仓库中。数据集、模型权重、视频和完整训练目录不直接放入 Git，而是通过 Release 资产提供下载链接。这样既能保持仓库体积可控，也便于从公开仓库重新检查代码结构和实验结论。",
    )
    add_table(
        doc,
        ["任务", "数据集", "模型", "核心指标", "当前最好结果"],
        [
            ["任务一", "Flowers102", "ResNet-18 / SE-ResNet-18", "Val/Test Accuracy", "Test Acc 0.8676"],
            ["任务二", "Road Vehicle", "YOLOv8n + BoT-SORT", "mAP / Tracking ID / Count", "mAP50 0.3770，计数 9"],
            ["任务三", "Stanford Background", "U-Net from scratch", "Pixel Acc / mIoU", "Val mIoU 0.5394"],
        ],
        "表2 三个任务的目标与结果概览",
    )
    add_picture_if_exists(doc, "reports/figures/flow_overall_hw2_gpt55.png", "图1 GPT-5.5 设计的 HW2 三任务总流程图", width_cm=15.5)

    add_heading(doc, "二、数据来源与预处理")
    add_paragraph(
        doc,
        "任务一使用 torchvision 提供的 Flowers102 标准划分。该数据集包含 102 类花卉，类别之间往往只在花瓣形状、颜色纹理或局部结构上有细微差异，因此非常适合检验 ImageNet 预训练特征在细粒度识别中的迁移能力。训练阶段将图像缩放到 224×224，并使用随机裁剪、水平翻转和 ImageNet 均值方差归一化；验证和测试阶段使用确定性缩放与中心裁剪，避免评估结果受到随机增强影响。",
    )
    add_paragraph(
        doc,
        "任务二使用 Road Vehicle Images Dataset。原始标注格式为 Supervisely rectangle JSON，代码先将标注解析为 YOLOv8 所需的 txt 文件，并生成 `data.yaml`。转换过程保留训练集 2704 张和验证集 300 张，共 21 个车辆相关类别。由于本地没有额外实拍道路视频，当前演示视频由验证集图像合成，主要用于稳定展示检测、跟踪、遮挡片段和越线计数流程。",
    )
    add_paragraph(
        doc,
        "任务三使用 Stanford Background Dataset。该数据集包含 715 张自然场景图像和像素级标签，类别包括 sky、tree、road、grass、water、building、mountain 和 foreground。代码固定随机种子后按 80%/20% 划分训练和验证集合，图像统一调整到 240×320，并对标签中的未知区域设置 ignore index，使其不参与损失和 mIoU 计算。",
    )
    add_table(
        doc,
        ["数据集", "规模/划分", "输入尺寸", "预处理", "评价方式"],
        [
            ["Flowers102", "标准 train/val/test", "224×224", "随机增强、ImageNet 归一化", "Accuracy"],
            ["Road Vehicle", "2704 train / 300 valid", "640", "JSON 转 YOLO 标注", "mAP、跟踪计数"],
            ["Stanford Background", "80% train / 20% valid", "240×320", "图像缩放、ignore index", "Pixel Acc、mIoU"],
        ],
        "表3 数据集与预处理方式",
    )

    add_heading(doc, "三、任务一：Flowers102 图像分类")
    add_picture_if_exists(doc, "reports/figures/flow_task1_flowers_gpt55.png", "图2 GPT-5.5 设计的 Flowers102 分类流程图", width_cm=15.5)
    add_heading(doc, "3.1 模型结构与训练策略", level=2)
    add_paragraph(
        doc,
        "分类 baseline 采用 ResNet-18。模型加载 ImageNet 预训练权重后，将原始 1000 类全连接分类头替换为 102 类输出层。为了避免大幅破坏预训练特征，优化器对骨干网络和分类头设置不同学习率：骨干网络使用较小学习率，分类头使用较大学习率。该策略符合迁移学习中的常见做法，即保留通用边缘、纹理和局部形状特征，同时让新分类头快速适配花卉类别。",
    )
    add_paragraph(
        doc,
        "注意力版本在 ResNet-18 最后一层卷积特征后加入 SE-block。SE-block 先通过全局平均池化得到通道描述向量，再使用两层全连接映射学习通道权重，最后将权重乘回原始特征。这样模型可以对更有判别力的花瓣纹理、颜色组合或局部结构赋予更高权重。由于 SE 模块参数量较小，它适合作为轻量注意力消融，不会显著改变 ResNet 主体结构。",
    )
    add_table(
        doc,
        ["实验项", "设置"],
        [
            ["基础模型", "ResNet-18"],
            ["输出类别", "102"],
            ["优化器", "AdamW"],
            ["学习率调度", "CosineAnnealingLR"],
            ["Batch size", "64"],
            ["权重衰减", "1e-4"],
            ["当前报告轮数", "8 epoch"],
        ],
        "表4 Flowers102 训练设置",
    )

    add_heading(doc, "3.2 超参数、预训练与注意力对比", level=2)
    add_paragraph(
        doc,
        "本次先进行 8 epoch 对比实验。虽然配置文件保留了 30 epoch 的完整训练命令，但短轮次已经能展示主要差异：预训练模型收敛速度明显快于随机初始化，学习率过小会拖慢分类头适配速度，而 SE 模块在 baseline 基础上带来小幅提升。对于 Flowers102 这样的细粒度数据集，小样本训练阶段的初始化质量和学习率设置会直接影响最终性能。",
    )
    add_table(
        doc,
        ["模型", "初始化", "注意力", "学习率", "最佳 Epoch", "Val Acc", "Test Acc"],
        [
            ["ResNet-18", "ImageNet", "无", "backbone 1e-4, head 1e-3", "7", "0.8873", "0.8653"],
            ["ResNet-18", "ImageNet", "无", "backbone 3e-5, head 3e-4", "7", "0.7108", "0.6720"],
            ["ResNet-18", "随机初始化", "无", "1e-3", "8", "0.2990", "0.2340"],
            ["ResNet-18", "ImageNet", "SE-block", "backbone 1e-4, head 1e-3", "7", "0.8980", "0.8676"],
        ],
        "表5 Flowers102 分类结果",
    )
    add_picture_if_exists(doc, "reports/figures/task1_summary.png", "图3 Flowers102 四组实验训练曲线汇总", width_cm=15.5)
    add_picture_if_exists(doc, "reports/figures/task1_resnet18_pretrained.png", "图4 ResNet-18 预训练 baseline 训练曲线", width_cm=15.5)
    add_picture_if_exists(doc, "reports/figures/task1_resnet18_random.png", "图5 ResNet-18 随机初始化消融曲线", width_cm=15.5)
    add_picture_if_exists(doc, "reports/figures/task1_resnet18_se.png", "图6 ResNet-18 + SE-block 训练曲线", width_cm=15.5)
    add_paragraph(
        doc,
        "从图1至图4可以看出，ImageNet 初始化带来的优势不仅体现在最终准确率，也体现在早期训练的稳定性。随机初始化曲线提升较慢，说明模型需要先从花卉数据中学习通用边缘和纹理，再学习类别判别；但 Flowers102 训练集规模有限，8 epoch 内很难充分完成这一过程。低学习率预训练组虽然同样使用 ImageNet 权重，但分类头学习速度不足，导致验证和测试准确率明显低于 baseline。SE-block 版本在验证集达到 0.8980，说明通道重标定对花卉细粒度纹理有一定帮助。",
    )
    add_paragraph(
        doc,
        "当前结果也说明，迁移学习实验不能只报告一个最优模型，还需要给出对照项。预训练对比回答了初始化是否重要，学习率对比说明优化策略是否合理，注意力对比则检验结构改动是否有效。三类对照共同支撑了最终选择 ResNet-18 + ImageNet + SE 作为当前最佳分类模型。",
    )

    add_heading(doc, "四、任务二：Road Vehicle 检测、跟踪与计数")
    add_picture_if_exists(doc, "reports/figures/flow_task2_tracking_gpt55.png", "图7 GPT-5.5 设计的道路车辆检测、跟踪与计数流程图", width_cm=15.5)
    add_heading(doc, "4.1 YOLOv8 检测模型", level=2)
    add_paragraph(
        doc,
        "目标检测部分采用 YOLOv8n。YOLOv8n 是较小规模的一阶段检测器，训练和推理速度较快，适合课程实验中快速验证完整流程。模型输入尺寸设为 640，训练使用 batch size 16。当前报告采用 20 epoch 结果，配置文件中保留 80 epoch 作为进一步提升 mAP 的完整设置。",
    )
    add_paragraph(
        doc,
        "Road Vehicle 数据集类别数较多，车辆姿态、遮挡、尺度和背景复杂度差异较大，因此 20 epoch 的 mAP50 仍有提升空间。当前训练重点是完成从数据格式转换、检测器微调、验证集评估到视频跟踪计数的端到端流程，并通过图表和日志说明模型行为。",
    )
    add_table(
        doc,
        ["项目", "设置"],
        [
            ["检测器", "YOLOv8n"],
            ["训练轮数", "20 epoch"],
            ["完整配置", "80 epoch"],
            ["输入尺寸", "640"],
            ["Batch size", "16"],
            ["验证指标", "Precision、Recall、mAP50、mAP50-95"],
        ],
        "表6 YOLOv8 检测训练设置",
    )
    add_table(
        doc,
        ["Epoch", "Precision", "Recall", "mAP50", "mAP50-95"],
        [["20", "0.6241", "0.3487", "0.3770", "0.2228"]],
        "表7 YOLOv8 验证结果",
    )
    add_picture_if_exists(doc, "reports/figures/task2_yolov8_results.png", "图8 YOLOv8 训练过程指标曲线", width_cm=15.5)
    add_picture_if_exists(doc, "reports/figures/task2_val_batch0_pred.jpg", "图9 Road Vehicle 验证集预测示例", width_cm=15.5)
    add_picture_if_exists(doc, "reports/figures/task2_confusion_matrix.png", "图10 YOLOv8 验证集混淆矩阵", width_cm=14.0)
    add_paragraph(
        doc,
        "图8展示了训练过程中的损失和 mAP 指标变化。随着 epoch 增加，模型逐步降低定位和分类损失，mAP50 在后期继续上升，说明 20 epoch 并未完全收敛。图9中的验证集预测可以直观看到模型已经能够定位多数车辆目标，但对密集目标、远距离小目标和相似类别仍会出现漏检或混淆。图10的混淆矩阵进一步说明，多类别车辆检测的主要困难来自类别外观相近和标注类别粒度较细。",
    )

    add_heading(doc, "4.2 多目标跟踪、遮挡分析与越线计数", level=2)
    add_paragraph(
        doc,
        "视频阶段使用 YOLOv8 内置 tracking 流程，并采用 BoT-SORT 进行跨帧目标关联。每一帧先由检测器输出车辆 bounding box、类别和置信度，再由跟踪器根据位置、运动和匹配分数维护 Tracking ID。计数逻辑不直接统计检测框数量，而是维护每个 ID 的上一帧中心点位置；当中心点从虚拟线一侧移动到另一侧时，且该 ID 尚未被计数，才累计一次。",
    )
    add_paragraph(
        doc,
        "当前演示视频为 12 秒合成片段，主要用于稳定复现多目标交汇和越线计数。跟踪日志共记录 2677 条检测结果、67 个唯一 Tracking ID，虚拟线设置为画面中线 `640,0,640,720`，最终得到 9 次越线计数。合成视频不是实拍道路视频，因此最终提交前如果有条件，可以替换为 10-30 秒真实道路片段，并复用同一脚本生成新日志和遮挡帧。",
    )
    add_table(
        doc,
        ["项目", "结果"],
        [
            ["视频长度", "12 秒"],
            ["跟踪日志行数", "2677"],
            ["唯一 Tracking ID", "67"],
            ["虚拟线", "640,0,640,720"],
            ["越线计数", "9"],
            ["输出视频", "runs/task2_detection/tracking_demo/tracked_counted.mp4"],
        ],
        "表8 跟踪与越线计数结果",
    )
    add_picture_if_exists(doc, "reports/figures/task2_overview_grid.jpg", "图11 检测、预测、混淆矩阵和遮挡片段 2×2 概览图", width_cm=15.5)
    add_picture_if_exists(doc, "reports/figures/task2_occlusion_grid.jpg", "图12 连续 4 帧遮挡与密集交汇片段", width_cm=15.5)
    add_paragraph(
        doc,
        "图12展示了连续 4 帧中多目标靠近虚拟线的过程。该片段的分析重点不是单帧检测框是否完全准确，而是同一目标在相邻帧中是否保持相同 ID。多数目标在连续帧中可以维持 ID，说明运动关联在短时间内较稳定；但当车辆框重叠、目标被部分遮挡或检测置信度降低时，跟踪器仍可能出现短暂丢失和新 ID。越线计数依赖 ID 连续性，因此 ID switch 会带来漏计或重复计风险。",
    )
    add_paragraph(
        doc,
        "为降低重复计数，代码对已经跨线的 ID 做去重处理，即同一 ID 只计一次。这比单纯统计某一侧检测框数量更稳定，也更符合交通计数场景。但该方法仍受检测质量影响：如果目标在跨线前后被漏检，中心点轨迹会断裂；如果目标类别混淆但 ID 保持稳定，计数仍然正确，但类别统计会受影响。后续若需要更高可靠性，可以进一步加入方向统计、轨迹平滑和最小轨迹长度过滤。",
    )

    add_heading(doc, "五、任务三：U-Net 语义分割与损失函数工程")
    add_picture_if_exists(doc, "reports/figures/flow_task3_unet_gpt55.png", "图13 GPT-5.5 设计的 U-Net 分割与损失对比流程图", width_cm=15.5)
    add_heading(doc, "5.1 U-Net 结构设计", level=2)
    add_paragraph(
        doc,
        "语义分割部分从零实现 U-Net，不加载任何预训练权重。模型由四级编码器、瓶颈层和四级解码器组成。编码器通过卷积块和下采样逐步扩大感受野，解码器通过上采样恢复空间分辨率，并与对应尺度的编码器特征做 skip connection。skip connection 能把浅层的边缘、纹理和位置信息传递到解码端，有助于恢复道路、树木和建筑边界。",
    )
    add_paragraph(
        doc,
        "每个卷积块使用两层 Conv-BN-ReLU。模型输出通道数等于 8 个语义类别，训练时对每个像素计算分类损失。由于 Stanford Background 的类别区域大小差异明显，单纯优化像素准确率容易偏向大面积类别，因此报告同时使用 mean IoU 衡量各类别区域重叠质量。",
    )
    add_table(
        doc,
        ["模块", "结构", "作用"],
        [
            ["Encoder", "4 级 Conv-BN-ReLU + 下采样", "提取多尺度上下文"],
            ["Bottleneck", "双卷积块", "聚合高层语义"],
            ["Decoder", "4 级上采样 + skip connection", "恢复空间细节"],
            ["Head", "1×1 卷积", "输出 8 类 logits"],
            ["初始化", "随机初始化", "满足从零训练"],
        ],
        "表9 U-Net 结构概览",
    )

    add_heading(doc, "5.2 Dice Loss 实现与损失对比", level=2)
    add_paragraph(
        doc,
        "本实验比较三种损失函数。Cross-Entropy 将每个像素视为独立多分类样本，训练稳定、梯度密集，是语义分割中常用的基础损失。Dice Loss 则从区域重叠角度出发，先对 logits 做 softmax，再构造 one-hot 标签，忽略未知像素后按类别计算 Dice 系数，最终使用 1-Dice 作为损失。CE+Dice 试图同时兼顾像素级分类和区域级重叠。",
    )
    add_paragraph(
        doc,
        "在 10 epoch 设置下，Cross-Entropy 的验证 mIoU 最高。原因可能有两点：第一，Stanford Background 的训练图像数量较少，Dice Loss 单独使用时早期优化不如 CE 稳定；第二，场景中 sky、tree、building 等大区域类别较多，CE 的像素级梯度更容易先学到全局类别分布。CE+Dice 的表现接近 CE，但没有超过 CE，说明在当前训练长度和模型规模下，组合损失还需要进一步调权或延长训练。",
    )
    add_table(
        doc,
        ["损失配置", "最佳 Epoch", "Val Pixel Acc", "Val mIoU"],
        [
            ["Cross-Entropy", "9", "0.7851", "0.5394"],
            ["Dice Loss", "9", "0.7484", "0.5033"],
            ["Cross-Entropy + Dice Loss", "9", "0.7775", "0.5308"],
        ],
        "表10 U-Net 分割结果",
    )
    add_picture_if_exists(doc, "reports/figures/task3_loss_grid.png", "图14 U-Net 三种损失配置 2×2 对比图", width_cm=15.5)
    add_picture_if_exists(doc, "reports/figures/task3_unet_ce.png", "图15 Cross-Entropy Loss 训练曲线", width_cm=15.5)
    add_picture_if_exists(doc, "reports/figures/task3_unet_dice.png", "图16 Dice Loss 训练曲线", width_cm=15.5)
    add_picture_if_exists(doc, "reports/figures/task3_unet_ce_dice.png", "图17 CE+Dice Loss 训练曲线", width_cm=15.5)
    add_paragraph(
        doc,
        "图14至图17显示，三种损失在训练后期均有提升，但收敛速度和稳定性不同。CE 曲线整体更平滑，验证 mIoU 也最高；Dice Loss 在区域重叠目标上更直接，但对类别分布和预测概率更敏感；CE+Dice 同时受到两种损失影响，验证性能介于二者之间。若继续训练到 50 epoch，可以观察 CE+Dice 是否在边界和小区域类别上逐渐超过单独 CE。",
    )

    add_heading(doc, "六、综合讨论")
    add_paragraph(
        doc,
        "三个任务共同说明，模型结构、训练策略和评价指标需要与空间智能问题的输出形式相匹配。分类任务只需输出图像类别，因此预训练特征和注意力模块是主要变量；检测跟踪任务需要同时关注位置和时间连续性，因此 mAP 之外还要分析 Tracking ID 稳定性；分割任务输出像素级标签，因此必须使用 mIoU 观察区域级质量，而不能只依赖 pixel accuracy。",
    )
    add_paragraph(
        doc,
        "从结果看，任务一效果最好，原因是 ImageNet 预训练提供了强先验，Flowers102 的视觉域也与自然图像接近。任务二 mAP 相对较低，主要受到类别粒度、目标尺度和训练轮数限制，但已经完成了从检测到跟踪计数的完整流程。任务三的 mIoU 处于可解释范围内，说明手写 U-Net 能学习到场景分割能力，但短轮次训练仍不足以充分处理细小区域和边界。",
    )
    add_paragraph(
        doc,
        "本报告保留了短轮次实验结果，是为了先形成完整、可检查的作业闭环。若继续优化，优先级可以按以下顺序展开：任务一训练到 30 epoch 并保存最优模型；任务二训练到 80 epoch 或更换更大 YOLOv8 模型，并使用真实道路视频；任务三使用 50 epoch 完整配置，并加入类别 IoU 表和预测可视化。这样可以在不改变代码结构的情况下进一步提高报告中的指标和分析深度。",
    )
    add_table(
        doc,
        ["方向", "当前处理", "后续改进"],
        [
            ["分类", "8 epoch 对比预训练、随机初始化、SE", "延长训练并尝试 ResNet-34"],
            ["检测", "YOLOv8n 20 epoch", "训练 80 epoch 或使用更大模型"],
            ["跟踪", "合成 12 秒演示视频", "替换为真实道路视频"],
            ["分割", "U-Net 10 epoch 三损失对比", "完整 50 epoch 并增加预测图"],
            ["报告", "Word 初稿可编辑", "补充小组信息并导出 PDF"],
        ],
        "表11 当前限制与改进方向",
    )

    add_heading(doc, "七、复现说明")
    add_paragraph(
        doc,
        "公开仓库中包含 `README.md`、`requirements.txt`、`configs/`、`src/`、`scripts/`、`reports/` 和 `outputs/`。复现时先安装依赖，再根据 README 中的命令分别运行三个任务。模型权重和跟踪演示视频不放入 Git 仓库，而是通过 Release 链接下载。报告中的精简结果日志位于 `outputs/logs/`，图表位于 `reports/figures/` 和 `outputs/figures/`。",
    )
    add_paragraph(
        doc,
        "若只检查报告结果，可直接查看 `reports/HW2_report_draft.docx`、`outputs/report/HW2_report_draft.docx` 和 `outputs/logs/`。若需要重新生成 Word 报告，可运行 `python -m src.common.make_docx_report --output reports/HW2_report_draft.docx`，再运行 `python scripts/collect_outputs.py` 同步到 `outputs/`。最终提交时建议在 Word 中填写小组成员、学号和分工，并手动导出 PDF。",
    )

    add_heading(doc, "八、结论")
    add_paragraph(
        doc,
        "本次 HW2 完成了分类、检测跟踪和语义分割三条实验线。Flowers102 实验表明，ImageNet 预训练和合适的学习率是细粒度分类的关键，SE 注意力带来小幅增益；Road Vehicle 实验完成了 YOLOv8 微调、验证集评估、Tracking ID 输出、遮挡片段分析和越线计数；Stanford Background 实验从零实现 U-Net，并验证了 CE、Dice 和 CE+Dice 三种损失的差异。",
    )
    add_paragraph(
        doc,
        "从作业目标看，当前代码已经覆盖模型结构、数据处理、训练设置、结果日志、可视化、公开仓库和权重链接。报告仍保留可编辑状态，便于继续加入更长训练结果、wandb 或 swanlab 截图、真实视频分析和最终小组信息。整体上，本实验展示了从图像级理解到目标级时序分析，再到像素级语义理解的完整空间智能建模流程。",
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate an editable Word report for HW2.")
    parser.add_argument("--output", default="reports/HW2_report_draft.docx")
    parser.add_argument("--template", default="")
    return parser.parse_args()


if __name__ == "__main__":
    args = parse_args()
    build_report(Path(args.output), Path(args.template) if args.template else None)
    print(args.output)
